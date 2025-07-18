#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os

def create_presentation():
    # Создаем новый документ
    doc = Document()
    
    # Настройка стилей
    styles = doc.styles
    
    # Стиль для заголовков
    heading_style = styles.add_style('CustomHeading', WD_STYLE_TYPE.PARAGRAPH)
    heading_style.font.name = 'Arial'
    heading_style.font.size = Pt(16)
    heading_style.font.bold = True
    heading_style.font.color.rgb = RGBColor(0, 0, 0)
    
    # Стиль для подзаголовков
    subheading_style = styles.add_style('CustomSubheading', WD_STYLE_TYPE.PARAGRAPH)
    subheading_style.font.name = 'Arial'
    subheading_style.font.size = Pt(14)
    subheading_style.font.bold = True
    subheading_style.font.color.rgb = RGBColor(44, 62, 80)
    
    # Основной заголовок
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('Sve Tu Marketplace')
    run.font.name = 'Arial'
    run.font.size = Pt(24)
    run.font.bold = True
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('Презентация для AtFrame DOO')
    run.font.name = 'Arial'
    run.font.size = Pt(18)
    
    doc.add_paragraph()
    
    # Вступление
    p = doc.add_paragraph('Добрый день, коллеги из AtFrame!')
    p.add_run('\n\nБлагодарим за готовность встретиться и поделиться опытом. Представляем краткую информацию о нашем проекте ')
    run = p.add_run('Sve Tu Marketplace')
    run.font.bold = True
    p.add_run('.')
    
    # О проекте
    doc.add_paragraph('О проекте', style='CustomHeading')
    
    p = doc.add_paragraph()
    run = p.add_run('Sve Tu')
    run.font.bold = True
    p.add_run(' (svetu.rs) - это современный онлайн-маркетплейс для Сербии, создаваемый как альтернатива существующим платформам с фокусом на:')
    
    doc.add_paragraph('• Безопасность сделок через эскроу-счета', style='List Bullet')
    doc.add_paragraph('• Удобный UX/UI для мобильных устройств', style='List Bullet')
    doc.add_paragraph('• Локальную адаптацию под сербский рынок', style='List Bullet')
    doc.add_paragraph('• Передовые технологии поиска и рекомендаций', style='List Bullet')
    
    # Наше решение
    doc.add_paragraph('Наше решение', style='CustomHeading')
    
    doc.add_paragraph('Для покупателей:', style='CustomSubheading')
    doc.add_paragraph('• Умный поиск с исправлением опечаток на сербском языке (fuzzy search с учетом кириллицы/латиницы)', style='List Bullet')
    doc.add_paragraph('• Интерактивная карта с пешим и радиусным поиском', style='List Bullet')
    doc.add_paragraph('• Защищенные сделки через эскроу-счета с интеграцией AllSecure/Stripe', style='List Bullet')
    doc.add_paragraph('• Real-time чат без обмена личными контактами, с вложениями и историей', style='List Bullet')
    doc.add_paragraph('• Избранное и подписки на категории с push-уведомлениями', style='List Bullet')
    doc.add_paragraph('• История просмотров и персональные рекомендации', style='List Bullet')
    doc.add_paragraph('• Система отзывов и рейтингов с фотографиями и модерацией', style='List Bullet')
    
    doc.add_paragraph('Для продавцов:', style='CustomSubheading')
    doc.add_paragraph('• Storefronts - полноценные интернет-магазины с кастомизацией', style='List Bullet')
    doc.add_paragraph('• Массовая загрузка товаров через CSV/XML/API с автоматическим маппингом категорий', style='List Bullet')
    doc.add_paragraph('• Варианты товаров (размеры, цвета) с управлением складом', style='List Bullet')
    doc.add_paragraph('• Детальная аналитика: просмотры, конверсии, воронки, heatmaps', style='List Bullet')
    doc.add_paragraph('• SEO-оптимизация: sitemap, schema.org, Open Graph', style='List Bullet')
    doc.add_paragraph('• История цен с графиками и уведомлениями покупателям', style='List Bullet')
    doc.add_paragraph('• Массовые операции: изменение цен, статусов, промо-акции', style='List Bullet')
    
    # Текущий статус
    doc.add_paragraph('Текущий статус', style='CustomHeading')
    doc.add_paragraph('✅ MVP запущен на svetu.rs', style='List Bullet')
    doc.add_paragraph('✅ 100+ тестовых объявлений', style='List Bullet')
    doc.add_paragraph('✅ Полная локализация (SR/RU/EN)', style='List Bullet')
    doc.add_paragraph('✅ Интеграция платежей (AllSecure/Stripe)', style='List Bullet')
    doc.add_paragraph('🔄 В разработке: варианты товаров, улучшение геолокации', style='List Bullet')
    
    # Безопасность и доверие
    doc.add_paragraph('Безопасность и доверие', style='CustomHeading')
    doc.add_paragraph('• Модерация контента: автоматическая проверка на спам, дубликаты, запрещенные товары', style='List Bullet')
    doc.add_paragraph('• Верификация продавцов: подтверждение личности, адреса, банковских данных', style='List Bullet')
    doc.add_paragraph('• Система доверия: рейтинги, бейджи, история сделок, отзывы с фото', style='List Bullet')
    doc.add_paragraph('• Защита от мошенничества: ML-алгоритмы выявления подозрительных паттернов', style='List Bullet')
    doc.add_paragraph('• Безопасные платежи: эскроу-счета, возврат средств, арбитраж споров', style='List Bullet')
    
    # Технические преимущества
    doc.add_paragraph('Технические преимущества', style='CustomHeading')
    
    doc.add_paragraph('Производительность:', style='CustomSubheading')
    doc.add_paragraph('• Оптимизация изображений: автоматическое сжатие, WebP, lazy loading', style='List Bullet')
    doc.add_paragraph('• Кеширование: Redis для горячих данных, CDN для статики', style='List Bullet')
    doc.add_paragraph('• База данных: PostgreSQL с PostGIS для гео-запросов', style='List Bullet')
    doc.add_paragraph('• Поиск: OpenSearch с fuzzy matching, синонимами, транслитерацией', style='List Bullet')
    
    doc.add_paragraph('Масштабируемость:', style='CustomSubheading')
    doc.add_paragraph('• Микросервисная архитектура: независимые модули', style='List Bullet')
    doc.add_paragraph('• Горизонтальное масштабирование: Docker Swarm ready', style='List Bullet')
    doc.add_paragraph('• API First: REST API для мобильных приложений и партнеров', style='List Bullet')
    doc.add_paragraph('• Webhooks: интеграция с внешними системами', style='List Bullet')
    
    # Новая страница для технологического стека
    doc.add_page_break()
    
    # Технологический стек
    doc.add_paragraph('Технологический стек', style='CustomHeading')
    
    doc.add_paragraph('Frontend:', style='CustomSubheading')
    doc.add_paragraph('• Next.js 15, React 19, TypeScript', style='List Bullet')
    doc.add_paragraph('• Tailwind CSS v4, DaisyUI', style='List Bullet')
    doc.add_paragraph('• Redux Toolkit, React Query', style='List Bullet')
    doc.add_paragraph('• Leaflet для интерактивных карт', style='List Bullet')
    
    doc.add_paragraph('Backend:', style='CustomSubheading')
    doc.add_paragraph('• Go (Golang) с Fiber framework', style='List Bullet')
    doc.add_paragraph('• PostgreSQL 15 + PostGIS', style='List Bullet')
    doc.add_paragraph('• OpenSearch для полнотекстового поиска', style='List Bullet')
    doc.add_paragraph('• MinIO (S3-compatible) для изображений', style='List Bullet')
    doc.add_paragraph('• Redis для кеширования', style='List Bullet')
    
    doc.add_paragraph('Инфраструктура:', style='CustomSubheading')
    doc.add_paragraph('• Docker + Docker Compose', style='List Bullet')
    doc.add_paragraph('• Nginx с Brotli сжатием', style='List Bullet')
    doc.add_paragraph('• GitHub Actions CI/CD', style='List Bullet')
    doc.add_paragraph('• Автоматические бэкапы', style='List Bullet')
    doc.add_paragraph('• Мониторинг: Prometheus + Grafana', style='List Bullet')
    
    # Бизнес-модель
    doc.add_paragraph('Бизнес-модель', style='CustomHeading')
    doc.add_paragraph('1. Комиссия 2-5% с успешных сделок через платформу', style='List Number')
    doc.add_paragraph('2. Платное продвижение объявлений (топ, выделение цветом)', style='List Number')
    doc.add_paragraph('3. Premium Storefronts с расширенными функциями:', style='List Number')
    doc.add_paragraph('   • Кастомный дизайн', style='List Bullet 2')
    doc.add_paragraph('   • Приоритет в поиске', style='List Bullet 2')
    doc.add_paragraph('   • Детальная аналитика', style='List Bullet 2')
    doc.add_paragraph('   • API доступ', style='List Bullet 2')
    doc.add_paragraph('   • Массовые операции', style='List Bullet 2')
    doc.add_paragraph('4. Дополнительные услуги:', style='List Number')
    doc.add_paragraph('   • Управление магазином', style='List Bullet 2')
    doc.add_paragraph('5. B2B решения: white-label, API для крупных ритейлеров', style='List Number')
    
    # Метрики и KPI
    doc.add_paragraph('Метрики и KPI', style='CustomHeading')
    
    doc.add_paragraph('Текущие показатели MVP:', style='CustomSubheading')
    doc.add_paragraph('• Время загрузки страницы: < 1.5 сек', style='List Bullet')
    doc.add_paragraph('• Uptime: 99.9%', style='List Bullet')
    doc.add_paragraph('• Конверсия просмотр → контакт: 3.5%', style='List Bullet')
    doc.add_paragraph('• Среднее время ответа API: 50ms', style='List Bullet')
    
    doc.add_paragraph('Целевые метрики на год:', style='CustomSubheading')
    doc.add_paragraph('• 50,000 активных объявлений', style='List Bullet')
    doc.add_paragraph('• 10,000 MAU (Monthly Active Users)', style='List Bullet')
    doc.add_paragraph('• 500 активных storefronts', style='List Bullet')
    doc.add_paragraph('• 1,000 транзакций в месяц', style='List Bullet')
    
    # Новая страница для планов и вопросов
    doc.add_page_break()
    
    # Планы развития
    doc.add_paragraph('Планы развития', style='CustomHeading')
    
    doc.add_paragraph('Краткосрочные (3-6 месяцев):', style='CustomSubheading')
    doc.add_paragraph('• Мобильные приложения (iOS/Android)', style='List Bullet')
    doc.add_paragraph('• Интеграция с местными службами доставки', style='List Bullet')
    doc.add_paragraph('• Расширенная аналитика для продавцов', style='List Bullet')
    doc.add_paragraph('• A/B тестирование и персонализация', style='List Bullet')
    
    doc.add_paragraph('Среднесрочные (6-12 месяцев):', style='CustomSubheading')
    doc.add_paragraph('• Экспансия в соседние страны (Черногория, Босния)', style='List Bullet')
    doc.add_paragraph('• B2B marketplace для оптовых покупателей', style='List Bullet')
    doc.add_paragraph('• AI-рекомендации и динамическое ценообразование', style='List Bullet')
    doc.add_paragraph('• Интеграция с ERP системами', style='List Bullet')
    
    # Что хотим обсудить
    doc.add_paragraph('Что хотим обсудить на встрече', style='CustomHeading')
    
    doc.add_paragraph('1. Технические аспекты:', style='CustomSubheading')
    doc.add_paragraph('• Интеграция с местными платежными системами (allsecure, payspot)', style='List Bullet')
    doc.add_paragraph('• Логистика (dexpress, posta srb)', style='List Bullet')
    doc.add_paragraph('• Опыт работы с местными хостинг-провайдерами', style='List Bullet')
    
    doc.add_paragraph('2. Юридические вопросы:', style='CustomSubheading')
    doc.add_paragraph('• GDPR vs сербское законодательство о защите данных', style='List Bullet')
    doc.add_paragraph('• Налоговая оптимизация для IT компаний', style='List Bullet')
    doc.add_paragraph('• Трудовые договоры и найм разработчиков', style='List Bullet')
    doc.add_paragraph('• Прочие юридические нюансы', style='List Bullet')
    
    doc.add_paragraph('3. Маркетинг и продвижение:', style='CustomSubheading')
    doc.add_paragraph('• Специфика digital-маркетинга в Сербии', style='List Bullet')
    doc.add_paragraph('• Опыт привлечения первых 1000 пользователей', style='List Bullet')
    doc.add_paragraph('• Работа с локальными инфлюенсерами', style='List Bullet')
    doc.add_paragraph('• Партнерства с традиционным бизнесом', style='List Bullet')
    doc.add_paragraph('• PR в местных СМИ', style='List Bullet')
    
    doc.add_paragraph('4. Бизнес-развитие:', style='CustomSubheading')
    doc.add_paragraph('• Особенности работы с сербскими инвесторами', style='List Bullet')
    doc.add_paragraph('• Гранты и поддержка от государства', style='List Bullet')
    doc.add_paragraph('• Опыт работы с Научно-технологическим парком', style='List Bullet')
    doc.add_paragraph('• Культурные особенности ведения бизнеса', style='List Bullet')
    doc.add_paragraph('• Конкуренция с устоявшимися игроками и возможные последствия', style='List Bullet')
    
    doc.add_paragraph('5. "Грабли" и лайфхаки:', style='CustomSubheading')
    doc.add_paragraph('• Типичные ошибки иностранных стартапов', style='List Bullet')
    doc.add_paragraph('• Неочевидные расходы и риски', style='List Bullet')
    doc.add_paragraph('• Полезные связи и организации', style='List Bullet')
    doc.add_paragraph('• Особенности менталитета пользователей', style='List Bullet')
    doc.add_paragraph('• Сезонность и локальные праздники', style='List Bullet')
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Подпись
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run('С уважением,\nКоманда Sve Tu Marketplace')
    
    # Сохраняем документ
    output_path = '/data/hostel-booking-system/Sve_Tu_Presentation_AtFrame.docx'
    doc.save(output_path)
    print(f"Документ успешно создан: {output_path}")

if __name__ == "__main__":
    create_presentation()