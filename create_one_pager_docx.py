#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

# Создаем новый документ
doc = Document()

# Настройка полей для one-pager
sections = doc.sections
for section in sections:
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

# Заголовок
title = doc.add_heading('SVE-TU', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.runs[0].font.size = Pt(36)
title.runs[0].font.color.rgb = RGBColor(102, 126, 234)

subtitle = doc.add_paragraph('AI-Powered Marketplace Platform')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].font.size = Pt(14)
subtitle.runs[0].font.bold = True

tagline = doc.add_paragraph('Революция e-commerce на Балканах')
tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
tagline.runs[0].font.size = Pt(12)
tagline.runs[0].font.italic = True

# Линия разделитель
doc.add_paragraph('_' * 50)

# Проблема и решение
doc.add_heading('ПРОБЛЕМА & РЕШЕНИЕ', 2)
problem = doc.add_paragraph()
problem.add_run('Проблема: ').bold = True
problem.add_run('50+ млн населения Балкан используют устаревшие маркетплейсы без AI, интеграций доставки и современного UX')

solution = doc.add_paragraph()
solution.add_run('Решение: ').bold = True
solution.add_run('Sve-Tu - первый на Балканах маркетплейс с полным циклом: от AI-создания объявлений до доставки и эскроу платежей')

# Ключевые метрики
doc.add_heading('КЛЮЧЕВЫЕ МЕТРИКИ', 2)

# Создаем таблицу для метрик
metrics_table = doc.add_table(rows=2, cols=3)
metrics_table.style = 'Light Grid Accent 1'

# Первая строка метрик
metrics_table.cell(0, 0).text = '85-90%\nMVP готов'
metrics_table.cell(0, 1).text = '10x\nэкономия на разработке'
metrics_table.cell(0, 2).text = '500K+\nстрок кода'

# Вторая строка метрик
metrics_table.cell(1, 0).text = '150+\nAPI endpoints'
metrics_table.cell(1, 1).text = '3 языка\nиз коробки'
metrics_table.cell(1, 2).text = '24/7\nAI разработка'

# Центрируем текст в таблице
for row in metrics_table.rows:
    for cell in row.cells:
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

# Уникальные преимущества
doc.add_heading('УНИКАЛЬНЫЕ ПРЕИМУЩЕСТВА', 2)

advantages = [
    ('AI-First подход', ['Создание объявления по 1 фото', 'Модерация в 100x быстрее', 'Персонализация +45% конверсии']),
    ('Claude Master System', ['-85% затрат на разработку', '5-7x скорость development', '24/7 автоматическая разработка']),
    ('Полная экосистема', ['C2C частные объявления', 'B2C/B2B корпоративные витрины', 'Собственные склады и ПВЗ', 'Fulfillment полный цикл', 'Эскроу платежи'])
]

for title, items in advantages:
    p = doc.add_paragraph()
    p.add_run(f'• {title}: ').bold = True
    p.add_run(', '.join(items))

# Бизнес-модель
doc.add_heading('БИЗНЕС-МОДЕЛЬ', 2)

bm_table = doc.add_table(rows=6, cols=3)
bm_table.style = 'Light List Accent 1'

# Заголовки
bm_table.cell(0, 0).text = 'Источник'
bm_table.cell(0, 1).text = 'Модель'
bm_table.cell(0, 2).text = 'Потенциал/год'

# Данные
bm_data = [
    ('Комиссия с продаж', '2-5% от транзакций', '€2-5M'),
    ('Fulfillment услуги', '€1-3/заказ + хранение', '€2-4M'),
    ('Premium подписки', '€29-299/мес', '€1-2M'),
    ('AI API & SaaS', 'Claude Master лицензии', '€3-10M'),
    ('Data Insights', 'Аналитика для брендов', '€0.5-1M')
]

for i, (source, model, potential) in enumerate(bm_data, 1):
    bm_table.cell(i, 0).text = source
    bm_table.cell(i, 1).text = model
    bm_table.cell(i, 2).text = potential

# Рынок и рост
doc.add_heading('РЫНОК & РОСТ', 2)

growth = doc.add_paragraph()
growth.add_run('2025: ').bold = True
growth.add_run('Сербия (7M) → €10M GMV\n')
growth.add_run('2026: ').bold = True
growth.add_run('Балканы (50M) → €100M GMV\n')
growth.add_run('2027: ').bold = True
growth.add_run('CEE (150M+) → €500M+ GMV')

# Конкурентные барьеры
doc.add_heading('КОНКУРЕНТНЫЕ БАРЬЕРЫ', 2)

barriers = [
    'Собственная логистика - склады + ПВЗ = контроль всей цепочки',
    'Технологический отрыв - Claude Master System недоступна конкурентам',
    'Локализация - 3 языка, локальные платежи и доставка',
    'AI интеграции - 12+ AI систем работают 24/7',
    'Скорость инноваций - новые функции за дни, не месяцы'
]

for barrier in barriers:
    doc.add_paragraph(f'✓ {barrier}')

# Инвестиционный раунд
doc.add_heading('ИНВЕСТИЦИОННЫЙ РАУНД', 2)

invest = doc.add_paragraph()
invest.add_run('Ищем: €1.5M Seed Round\n').bold = True
invest.add_run('• Маркетинг & User Acquisition - 40%\n')
invest.add_run('• Развитие продукта & AI - 30%\n')
invest.add_run('• Операции & Логистика - 20%\n')
invest.add_run('• Команда - 10%\n\n')
invest.add_run('Предлагаем: ').bold = True
invest.add_run('15-20% equity | Board seat | Exit 3-5 лет')

# Roadmap
doc.add_heading('ROADMAP', 2)

roadmap = [
    ('Q1 2025', 'Запуск в Сербии, 1K продавцов'),
    ('Q2 2025', '10K листингов, break-even'),
    ('Q3 2025', 'Экспансия: Хорватия, Босния'),
    ('Q4 2025', 'Series A, вся ex-Югославия'),
    ('2026', 'Лидер Балканского e-commerce')
]

for quarter, milestone in roadmap:
    p = doc.add_paragraph()
    p.add_run(f'{quarter}: ').bold = True
    p.add_run(milestone)

# Почему сейчас
doc.add_heading('ПОЧЕМУ СЕЙЧАС?', 2)

reasons = [
    'AI революция - первые получают весь рынок',
    'Растущий e-commerce - +35% год к году на Балканах',
    'Готовый продукт - не концепт, а работающая платформа',
    'Уникальная технология - Claude Master = несправедливое преимущество'
]

for i, reason in enumerate(reasons, 1):
    doc.add_paragraph(f'{i}. {reason}')

# Контакты
doc.add_heading('КОНТАКТЫ', 2)

contacts = doc.add_paragraph()
contacts.add_run('Email: ').bold = True
contacts.add_run('investors@svetu.rs\n')
contacts.add_run('Demo: ').bold = True
contacts.add_run('https://svetu.rs\n')
contacts.add_run('Deck: ').bold = True
contacts.add_run('По запросу (NDA)\n\n')

# Финальный призыв
final_cta = doc.add_paragraph()
final_cta.add_run('Присоединяйтесь к революции e-commerce на Балканах!\n').bold = True
final_cta.alignment = WD_ALIGN_PARAGRAPH.CENTER

slogan = doc.add_paragraph()
slogan.add_run('Sve-Tu - "Все здесь" на сербском. Мы создаем место, где есть всё.').italic = True
slogan.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Сохраняем документ
doc.save('/data/hostel-booking-system/SVE_TU_ONE_PAGER.docx')

print('One-pager DOCX успешно создан: SVE_TU_ONE_PAGER.docx')