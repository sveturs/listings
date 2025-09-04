from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import io

# Создаем документ
doc = Document()

# Настраиваем стили
style = doc.styles['Normal']
style.font.name = 'Arial'
style.font.size = Pt(11)

# Добавляем тему
title = doc.add_paragraph()
title_run = title.add_run('Тема: Предложение по реализации безопасных C2C сделок - Sve Tu Platforma')
title_run.bold = True
title_run.font.size = Pt(12)

# Пустая строка
doc.add_paragraph()

# Приветствие
doc.add_paragraph('Поштована Снежана,\nПоштована Тамара,')
doc.add_paragraph()

# Основной текст
doc.add_paragraph('Благодарю за продуктивную встречу. Продолжим обсуждение перспектив развития C2C сегмента. В соответствии с нашей договоренностью, направляю вам детализированное предложение по созданию системы безопасных сделок между физическими лицами, которая позволит существенно снизить уровень мошенничества на рынке подержанных товаров Сербии.')

doc.add_paragraph('Прежде всего, хочу подчеркнуть, что поиск возможности реализации escrow C2C в Сербии не должен тормозить процесс подписания договора с ранее согласованными параметрами для рынка B2C. Мы готовы двигаться параллельно по обоим направлениям.')

doc.add_paragraph('Изучив успешный опыт реализации защиты сделок на других платформах за пределами Сербии (Avito в России, OLX Pay в Польше, Vinted в Европе), мы подготовили комплексное решение, которое полностью адресует озвученные вами опасения относительно AML compliance и возможного использования C2C для обхода налогов.')

# Заголовок раздела
heading1 = doc.add_paragraph()
heading1_run = heading1.add_run('ТЕХНИЧЕСКАЯ СХЕМА ДВИЖЕНИЯ СРЕДСТВ C2C')
heading1_run.bold = True
heading1_run.font.size = Pt(14)
heading1.alignment = WD_ALIGN_PARAGRAPH.LEFT

p = doc.add_paragraph()
p_run = p.add_run('Важно:')
p_run.bold = True
p.add_run(' В отличие от B2C, где AllSecure обеспечивает временную блокировку средств инструментами Visa/Mastercard, для C2C предлагаем следующую схему через PaySpot:')

# Нумерованный список
doc.add_paragraph('1. ', style='List Number').add_run('Оплата покупателем:').bold = True
doc.add_paragraph('   - Покупатель оплачивает через PaySpot (карта/банковский перевод)')
doc.add_paragraph('   - Средства поступают на транзитный счет PaySpot')
doc.add_paragraph('   - Статус "Ожидает отправки" в системе')

doc.add_paragraph('2. ', style='List Number').add_run('Период безопасности (5-14 дней):').bold = True
doc.add_paragraph('   - Продавец отправляет товар и вводит трек-номер')
doc.add_paragraph('   - Покупатель получает и проверяет товар')
doc.add_paragraph('   - Подтверждает получение или открывает спор')

doc.add_paragraph('3. ', style='List Number').add_run('Выплата продавцу - физическому лицу:').bold = True
doc.add_paragraph('   - Вариант А: На PaySpot Wallet продавца (с возможностью вывода на карту)').add_run(' Вариант А:').bold = True
doc.add_paragraph('   - Вариант Б: Прямой перевод на банковский счет через IPS (instant payment)').add_run(' Вариант Б:').bold = True
doc.add_paragraph('   - Вариант В: На привязанную дебетовую карту продавца').add_run(' Вариант В:').bold = True
doc.add_paragraph()
doc.add_paragraph('   Комиссия PaySpot взимается автоматически перед выплатой.')

# Следующий раздел
heading2 = doc.add_paragraph()
heading2_run = heading2.add_run('КЛЮЧЕВЫЕ ЭЛЕМЕНТЫ ЗАЩИТЫ:')
heading2_run.bold = True
heading2_run.font.size = Pt(14)

# Подраздел 1
subheading1 = doc.add_paragraph()
subheading1_run = subheading1.add_run('1. Защита от отмывания денег (AML)')
subheading1_run.bold = True
subheading1_run.font.size = Pt(12)

p = doc.add_paragraph()
p_run = p.add_run('Многоуровневая система проверок:')
p_run.bold = True

doc.add_paragraph('- Пороговые значения с автоматической верификацией (SMS при сделках >100 EUR, фото паспорта >300 EUR, видеоверификация >500 EUR)')
doc.add_paragraph('- AML скоринг каждой транзакции с анализом паттернов поведения')
doc.add_paragraph('- Интеграция с государственными базами (APR, Централни регистар фактура)')
doc.add_paragraph('- Ежемесячная отчетность для PaySpot по всем транзакциям свыше 1,000 EUR')

# Подраздел 2
subheading2 = doc.add_paragraph()
subheading2_run = subheading2.add_run('2. Предотвращение обхода налогов')
subheading2_run.bold = True
subheading2_run.font.size = Pt(12)

p = doc.add_paragraph()
p_run = p.add_run('Четкое разделение C2C и B2C:')
p_run.bold = True

doc.add_paragraph('- AI-анализ фотографий для определения новых/БУ товаров')
doc.add_paragraph('- Автоматическое выявление признаков коммерческой деятельности (>5 одинаковых товаров, профессиональные фото, систематические продажи)')
doc.add_paragraph('- Жесткие лимиты: максимум 10 активных объявлений, до 2,000 EUR/месяц')
doc.add_paragraph('- При превышении лимитов - автоматическое перенаправление в B2C с полной проверкой документов')

# Подраздел 3
subheading3 = doc.add_paragraph()
subheading3_run = subheading3.add_run('3. Прозрачность и отчетность')
subheading3_run.bold = True
subheading3_run.font.size = Pt(12)

doc.add_paragraph('- Ежегодные справки продавцам о суммах транзакций')
doc.add_paragraph('- Готовность к сотрудничеству с Пореском управом')
doc.add_paragraph('- Полное соблюдение требований НБС')

# Предложение пилотного проекта
heading3 = doc.add_paragraph()
heading3_run = heading3.add_run('ПРЕДЛОЖЕНИЕ ПИЛОТНОГО ПРОЕКТА')
heading3_run.bold = True
heading3_run.font.size = Pt(14)

doc.add_paragraph('Предлагаем запустить 3-месячный пилот с ограничениями:')
doc.add_paragraph('- Только БУ товары в категориях: электроника, одежда, книги')
doc.add_paragraph('- Максимум 300 EUR за транзакцию')
doc.add_paragraph('- Максимум 1,000 EUR/месяц на пользователя')
doc.add_paragraph('- 5-50 проверенных пользователей с полной верификацией')

p = doc.add_paragraph()
p_run = p.add_run('Ожидаемые результаты пилота:')
p_run.bold = True

doc.add_paragraph('- 0 инцидентов, связанных с AML')
doc.add_paragraph('- <1% попыток продажи новых товаров через C2C')
doc.add_paragraph('- Полная аналитика для дальнейшего масштабирования')

# Выгоды для PaySpot
heading4 = doc.add_paragraph()
heading4_run = heading4.add_run('ВЫГОДЫ ДЛЯ PAYSPOT')
heading4_run.bold = True
heading4_run.font.size = Pt(14)

doc.add_paragraph('1. Лидерство на рынке - первое полностью compliant C2C решение в Сербии').add_run('Лидерство на рынке').bold = True
doc.add_paragraph('2. Дополнительный revenue stream без увеличения рисков').add_run('Дополнительный revenue stream').bold = True
doc.add_paragraph('3. Ценные данные для улучшения скоринговых моделей').add_run('Ценные данные').bold = True
doc.add_paragraph('4. Возможность лицензирования решения другим платформам региона').add_run('Возможность лицензирования').bold = True

# Заключительная часть
doc.add_paragraph()
doc.add_paragraph('Мы полностью понимаем вашу ответственность перед регулятором и готовы работать в самых строгих рамках compliance. Наша цель - создать безопасную среду для всех участников рынка при полном соблюдении законодательства.')

doc.add_paragraph('В ответ жду обсуждение предложенного нами варианта, если у Вас есть вопросы или предложения вариантов по улучшению для предоставления возможности Escrow для C2C - давайте обсуждать.')

doc.add_paragraph()
doc.add_paragraph('С наилучшими пожеланиями,')
doc.add_paragraph()

# Подпись
signature = doc.add_paragraph()
signature_run1 = signature.add_run('Дмитрий Ворошилов')
signature_run1.bold = True
doc.add_paragraph('Директор')
doc.add_paragraph('Sve Tu Platforma DOO Novi Sad')
doc.add_paragraph('Email: dima@svetu.rs , docs@svetu.rs')

# P.S.
doc.add_paragraph()
ps = doc.add_paragraph('P.S. Еще раз подтверждаю нашу готовность подписать договор по B2C в ранее согласованных параметрах независимо от решения по C2C направлению.')

# Сохраняем документ
buffer = io.BytesIO()
doc.save(buffer)
buffer.seek(0)

# Сохраняем файл
with open('PaySpot_C2C_Proposal.docx', 'wb') as f:
    f.write(buffer.getvalue())

print("Документ 'PaySpot_C2C_Proposal.docx' успешно создан!")