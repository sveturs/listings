#!/usr/bin/env python3
"""
Script to generate DOCX documents for Infobip Viber Bot application
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from datetime import datetime
import os

def create_warranties_letter():
    """Create the BM Warranties Letter document"""
    doc = Document()

    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Company Letterhead
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('SVE TU PLATFORMA DOO\n')
    run.bold = True
    run.font.size = Pt(16)
    p.add_run('Vase Stajića 18, apartment 18\n21000 Novi Sad, Serbia\nPIB: 113569408\nMB: 21891975')

    doc.add_paragraph()
    doc.add_paragraph('─' * 50)
    doc.add_paragraph()

    # To line
    p = doc.add_paragraph()
    p.add_run('To: ').bold = True
    p.add_run('Viber Media S.a r.l ("Viber")')

    # Re line
    p = doc.add_paragraph()
    p.add_run('Re: ').bold = True
    p.add_run('Warranties regarding Sending Verified Messages via Viber')

    doc.add_paragraph()

    # Main text
    text = """The undersigned, SVE TU PLATFORMA DOO (the "Company"), owner of the SveTu Marketplace brand, whose ultimate Beneficial Owners are Dmitrii Voroshilov (50%) and Azamat Salakhov (50%) and whose headquarters are located at Vase Stajića 18, apartment 18, 21000 Novi Sad, Serbia as appears on Company's registration documents/business license, wishes to use Viber's Business Messages feature to send messages to its opted in customers."""

    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    doc.add_paragraph()
    doc.add_paragraph('In connection with such intended use, it hereby undertakes that as part of the sending of business messages through the verified messages channel by Viber, it will:')

    # Numbered list of obligations
    obligations = [
        'Send messages only to users who have consented to receive messages from it and did not revoke such consent, and who are at legal age to receive such messages according to applicable laws.',
        'Comply with all applicable laws in connection with the sending of the messages, and ensure that the texts are compliant with such laws.',
        'Send messages only in accordance with the Business Messages Guidelines provided by Viber.',
        'Will not allow any third party to use its verified messages channel.',
        'Not use the verified messages channel to promote, mention and/or send invites to, competitive platforms/messaging apps.'
    ]

    for i, obligation in enumerate(obligations, 1):
        p = doc.add_paragraph(f'{i}. {obligation}')
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    doc.add_paragraph()

    # Spam handling paragraph
    spam_text = """If Viber contacts the Company with regards to a complaint of a Spam incident, i.e. a message sent to a non-opted in user or a message which violates the Business Messages Guidelines ("Spam"), received by Viber support team, the Company shall: (i) respond as soon as possible, and in any case within no more than 7 days, (ii) provide Viber with all necessary information (opt-in documentation, relevant employee/department contact details, and any other relevant reasonable information required by Viber) within 14 days, and (iii) take such complaints very seriously and handle accordingly. The Company shall promptly provide Viber with several responsive contact methods to be provided to the relevant user to directly contact the Company with regards to such Spam event."""

    p = doc.add_paragraph(spam_text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    doc.add_paragraph()

    # Indemnification paragraph
    indemnity_text = """The Company agrees to indemnify, defend and hold harmless Viber, its officers, employees and affiliates in connection with any claim by a third party (including government entities) in connection with any message sent via its verified messages channel. It further agrees and consents that Viber shall have the right to reject and to terminate any service, at its sole discretion, without notice and for any reason."""

    p = doc.add_paragraph(indemnity_text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    doc.add_paragraph()

    # Liquidated damages
    damages_text = """Without derogating from the above and without prejudice to any other rights and remedies available to Viber, for every incident of Spam (as defined above), Company agrees that Viber shall charge agreed liquidated damages in the amount of €5,000."""

    p = doc.add_paragraph(damages_text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    doc.add_paragraph()

    # Governing law
    law_text = """This letter is governed by the laws of England and Wales and subject to the sole jurisdiction of the competent courts of London, UK."""

    p = doc.add_paragraph(law_text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    doc.add_paragraph()
    doc.add_paragraph('IN WITNESS WHEREOF,')
    doc.add_paragraph()
    doc.add_paragraph('The authorized signatory on behalf of the Company has confirmed the warranties and obligations included herein.')
    doc.add_paragraph()
    doc.add_paragraph()

    # Signature block
    doc.add_paragraph('Signature: _______________________')
    doc.add_paragraph()
    p = doc.add_paragraph('Name: ')
    p.add_run('Dmitrii Voroshilov').bold = True
    doc.add_paragraph()
    p = doc.add_paragraph('Title: ')
    p.add_run('Director').bold = True
    doc.add_paragraph()
    p = doc.add_paragraph('Date: ')
    p.add_run(datetime.now().strftime('%B %d, %Y')).bold = True

    # Save document
    doc.save('/data/hostel-booking-system/docs/Infobip_Warranties_Letter.docx')
    print("✅ Created: Infobip_Warranties_Letter.docx")

def create_qualification_form():
    """Create the Chatbot Qualification Form document"""
    doc = Document()

    # Title
    title = doc.add_heading('Chatbot Qualification Form', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # KEY FEATURE HIGHLIGHT
    highlight = doc.add_heading('🎯 KEY INNOVATION: Real-Time Courier Tracking Inside Viber', 1)
    highlight.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Revolutionary delivery tracking experience - users can see their courier moving on an interactive map ')
    run.bold = True
    run.font.size = Pt(12)
    run = p.add_run('DIRECTLY INSIDE VIBER')
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(128, 0, 128)  # Purple color for Viber

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run('No app switching • No external links • Live GPS tracking in Viber\'s embedded browser')

    doc.add_paragraph('─' * 50)
    doc.add_paragraph()

    # General Information section
    doc.add_heading('General Information', 1)

    # Create table
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'

    # Header row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Parameter'
    hdr_cells[1].text = 'Additional info'
    hdr_cells[2].text = "Partner's Value"

    # Make header bold
    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    # Data rows
    data = [
        ('Bot name', 'Visible on Viber\nUp to 28 alphanumeric chars', 'SveTu Marketplace'),
        ('Bot URI', 'Visible in the share links\nUp to 28 alphanumeric chars, no spaces allowed', 'svetumarketplace'),
        ('Bot Category', 'Visible on Viber\nOne of: Local Businesses, Companies, Brands & Products, People, Places of Interest, Groups, Organizations & Institutions, Entertainment', 'Companies'),
        ('Bot Description', 'Visible on Viber\nUp to 28 alphanumeric chars', 'Serbian marketplace platform'),
        ('Bot Language', '', 'Serbian, English, Russian'),
        ('Website Address', 'Visible on Viber', 'https://svetu.rs'),
        ("Bot owner's email address", 'For Viber Internal purposes', 'docs@svetu.rs'),
        ("Admin's phone number", 'Phone number of the viber user, who will own the bot and have access to its management.\nThe phone number must be registered in Viber.\nPlease enter full number, including country code with no +, - or leading zeros', '381629316318'),
        ('Launch date estimation', '', 'January 2025'),
        ('Bot location', 'Visible on Viber\nCountry and city for bot info page', 'Serbia, Novi Sad'),
        ('Commercial Account ID', 'Relevant if you want to connect a bot to a Commercial Account', '[To be provided by Infobip]')
    ]

    for param, info, value in data:
        row_cells = table.add_row().cells
        row_cells[0].text = param
        row_cells[1].text = info
        row_cells[2].text = value
        # Bold the values
        for paragraph in row_cells[2].paragraphs:
            for run in paragraph.runs:
                run.bold = True

    doc.add_paragraph()

    # Verification & Searchability section
    doc.add_heading('Verification & Searchability', 1)
    doc.add_paragraph('After the bot is created and ready for publishing, you can reach out to Viber team and request the settings')

    doc.add_paragraph()

    # Account content restriction
    p = doc.add_paragraph()
    p.add_run('Account content restriction: ').bold = True
    p.add_run('No restrictions')

    doc.add_paragraph()
    doc.add_paragraph('─' * 50)
    doc.add_paragraph()

    # Additional details
    doc.add_heading('Additional Details', 1)

    # Company information
    doc.add_heading('Company Information:', 2)
    info_items = [
        ('Legal Entity:', 'SVE TU PLATFORMA DOO'),
        ('Registration Number:', 'MB: 21891975'),
        ('Tax ID:', 'PIB: 113569408'),
        ('Address:', 'Vase Stajića 18, apartment 18, 21000 Novi Sad, Serbia'),
        ('Beneficial Owners:', 'Dmitrii Voroshilov (50%), Azamat Salakhov (50%)'),
        ('Directors:', 'Dmitrii Voroshilov, Azamat Salakhov')
    ]

    for label, value in info_items:
        p = doc.add_paragraph()
        p.add_run(f'{label} ').bold = True
        p.add_run(value)

    doc.add_paragraph()

    # Bot purpose
    doc.add_heading('Bot Purpose & Unique Value Proposition:', 2)

    # Main killer feature
    doc.add_heading('PRIMARY FEATURE - Real-Time Delivery Tracking:', 3)
    tracking_text = """When a user asks "Where is my package?" (Где моя посылка?), our bot provides:

1. INSTANT RESPONSE with Rich Media card containing:
   • Static map showing delivery route
   • Current courier location
   • Estimated time of arrival
   • "Open Live Map" button

2. REVOLUTIONARY LIVE TRACKING - When user clicks the button:
   • Opens interactive map INSIDE Viber's embedded browser
   • Shows courier moving in REAL-TIME using GPS
   • Updates location every 5-10 seconds via WebSocket
   • Displays speed, direction, and ETA
   • Works WITHOUT leaving Viber app
   • NO app installation required

This is a GAME-CHANGER for Serbian e-commerce - no other marketplace offers real-time courier tracking directly inside Viber!"""

    p = doc.add_paragraph(tracking_text)
    for paragraph in p.runs:
        if 'REAL-TIME' in paragraph.text or 'GAME-CHANGER' in paragraph.text:
            paragraph.bold = True

    doc.add_paragraph()

    doc.add_heading('Additional Features:', 3)
    purpose_text = """• Customer support and FAQs
• Product search functionality
• Order status notifications
• Notifications for sellers about new orders
• Shopping cart reminders
• Promotional messages (only to opted-in users)"""

    doc.add_paragraph(purpose_text)

    doc.add_paragraph()

    # Expected message volume
    doc.add_heading('Expected Message Volume:', 2)
    doc.add_paragraph('• Initial: 5,000-10,000 messages/month')
    doc.add_paragraph('• Growth target: 50,000+ messages/month within first year')

    doc.add_paragraph()

    # Signature
    doc.add_paragraph('─' * 50)
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Form prepared by: ').italic = True
    p.add_run('Dmitrii Voroshilov')

    p = doc.add_paragraph()
    p.add_run('Date: ').italic = True
    p.add_run(datetime.now().strftime('%B %d, %Y'))

    p = doc.add_paragraph()
    p.add_run('Company: ').italic = True
    p.add_run('SVE TU PLATFORMA DOO')

    # Save document
    doc.save('/data/hostel-booking-system/docs/Infobip_Chatbot_Qualification_Form.docx')
    print("✅ Created: Infobip_Chatbot_Qualification_Form.docx")

def main():
    """Main function to generate all documents"""
    print("📄 Generating DOCX documents for Infobip application...")
    print()

    try:
        # Check if python-docx is installed
        import docx
    except ImportError:
        print("❌ Error: python-docx is not installed")
        print("Please run: pip install python-docx")
        return

    # Create documents
    create_warranties_letter()
    create_qualification_form()

    print()
    print("✅ All documents have been created successfully!")
    print()
    print("📁 Files location: /data/hostel-booking-system/docs/")
    print("   • Infobip_Warranties_Letter.docx")
    print("   • Infobip_Chatbot_Qualification_Form.docx")
    print()
    print("📋 Next steps:")
    print("1. Review and print the Warranties Letter on company letterhead")
    print("2. Sign the Warranties Letter")
    print("3. Fill in any remaining details in the Qualification Form")
    print("4. Send both documents to Infobip")

if __name__ == "__main__":
    main()