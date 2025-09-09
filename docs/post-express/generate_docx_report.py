#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import os

def add_code_block(doc, code, language=""):
    """Add a formatted code block to the document"""
    p = doc.add_paragraph()
    p.style = 'Normal'
    run = p.add_run(code)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    # Add light gray background
    shading_elm = parse_xml(r'<w:shd {} w:fill="F0F0F0"/>'.format(nsdecls('w')))
    p._element.get_or_add_pPr().append(shading_elm)
    return p

def create_technical_report():
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)
    
    # Title
    title = doc.add_heading('TEHNICKI IZVESTAJ ZA PODRSKU POST EXPRESS', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Metadata
    doc.add_paragraph('Datum: 2025-09-08')
    doc.add_paragraph('Od: Sve Tu Platforma')
    doc.add_paragraph('Za: Tehnicku podrsku Post Express (b2b@posta.rs)')
    doc.add_paragraph('Tema: Problem sa pristupnim podacima TEST u WSP API')
    doc.add_paragraph()
    
    # Section 1: Rezime problema
    doc.add_heading('REZIME PROBLEMA', level=2)
    doc.add_paragraph(
        'Ne mozemo uspesno kreirati posiljku kroz WSP API koristeci pristupne podatke '
        'TEST/t3st iz vase dokumentacije. API vraca gresku: "Korisnicko ime TEST nije registrovano!"'
    )
    doc.add_paragraph()
    
    # Section 2: NASA IMPLEMENTACIJA
    doc.add_heading('NASA IMPLEMENTACIJA', level=2)
    
    doc.add_heading('1. Struktura zahteva (Go kod)', level=3)
    code1 = '''// backend/internal/proj/postexpress/models/models.go
type TransakcijaIn struct {
    StrKlijent         string `json:"StrKlijent"`
    Servis             int    `json:"Servis"`
    IdVrstaTranskacije int    `json:"IdVrstaTranskacije"` // Sa slovom "k"!
    TipSerijalizacije  int    `json:"TipSerijalizacije"`
    IdTransakcija      string `json:"IdTransakcija"`
    StrIn              string `json:"StrIn"`
}

type Klijent struct {
    Username      string `json:"Username"`
    Password      string `json:"Password"`
    Jezik         string `json:"Jezik"`
    IdTipUredjaja int    `json:"IdTipUredjaja"`
}'''
    add_code_block(doc, code1, "go")
    doc.add_paragraph()
    
    doc.add_heading('2. Funkcija slanja zahteva', level=3)
    code2 = '''// backend/internal/proj/postexpress/service/client.go
func (c *Client) SendRequest(req *TransakcijaIn) (*TransakcijaOut, error) {
    // Formiramo Klijent
    klijent := Klijent{
        Username:      c.username,  // "TEST"
        Password:      c.password,  // "t3st"
        Jezik:         "LAT",
        IdTipUredjaja: 2,
    }
    
    klijentJSON, _ := json.Marshal(klijent)
    req.StrKlijent = string(klijentJSON)
    
    // Saljemo zahtev
    jsonData, _ := json.Marshal(req)
    
    httpReq, _ := http.NewRequest("POST", 
        "http://212.62.32.201/WspWebApi/transakcija", 
        bytes.NewBuffer(jsonData))
    httpReq.Header.Set("Content-Type", "application/json")
    
    resp, _ := c.httpClient.Do(httpReq)
    // ... obrada odgovora
}'''
    add_code_block(doc, code2, "go")
    doc.add_paragraph()
    
    # Section 3: TEST ZAHTEVI I ODGOVORI
    doc.add_heading('TEST ZAHTEVI I ODGOVORI', level=2)
    
    doc.add_heading('Test 1: Prosta transakcija GetNaselje (ID=3)', level=3)
    doc.add_paragraph('Zahtev:')
    json1 = '''{
  "StrKlijent": "{\\"Username\\":\\"TEST\\",\\"Password\\":\\"t3st\\",\\"Jezik\\":\\"LAT\\",\\"IdTipUredjaja\\":2}",
  "Servis": 3,
  "IdVrstaTranskacije": 3,
  "TipSerijalizacije": 2,
  "IdTransakcija": "test-1736955123",
  "StrIn": "{\\"Naziv\\":\\"Novi\\"}"
}'''
    add_code_block(doc, json1, "json")
    
    doc.add_paragraph('cURL komanda:')
    curl1 = '''curl -X POST http://212.62.32.201/WspWebApi/transakcija \\
  -H "Content-Type: application/json" \\
  -d '{"StrKlijent":"{\\"Username\\":\\"TEST\\",\\"Password\\":\\"t3st\\",\\"Jezik\\":\\"LAT\\",\\"IdTipUredjaja\\":2}","Servis":3,"IdVrstaTranskacije":3,"TipSerijalizacije":2,"IdTransakcija":"test-1736955123","StrIn":"{\\"Naziv\\":\\"Novi\\"}"}'
'''
    add_code_block(doc, curl1, "bash")
    
    doc.add_paragraph('Odgovor:')
    response1 = '''{
  "Rezultat": 3,
  "StrOut": null,
  "StrRezultat": "{
    \\"Poruka\\": \\"Korisnicko ime TEST nije registrovano!\\",
    \\"PorukaKorisnik\\": \\"Korisnicko ime TEST nije registrovano!\\"
  }"
}'''
    add_code_block(doc, response1, "json")
    doc.add_paragraph()
    
    doc.add_heading('Test 2: Transakcija Manifest (ID=73) za kreiranje posiljke', level=3)
    doc.add_paragraph('Zahtev:')
    json2 = '''{
  "StrKlijent": "{\\"Username\\":\\"TEST\\",\\"Password\\":\\"t3st\\",\\"Jezik\\":\\"LAT\\",\\"IdTipUredjaja\\":2}",
  "Servis": 3,
  "IdVrstaTranskacije": 73,
  "TipSerijalizacije": 2,
  "IdTransakcija": "manifest-1736955456",
  "StrIn": "{
    \\"Posiljalac\\": {
      \\"Ime\\": \\"Test Sender\\",
      \\"Adresa\\": \\"Bulevar kralja Aleksandra 1\\",
      \\"IdNaselje\\": 110000,
      \\"Telefon\\": \\"0601234567\\"
    },
    \\"Posiljke\\": [{
      \\"Primalac\\": {
        \\"Ime\\": \\"Test Receiver\\",
        \\"Adresa\\": \\"Knez Mihailova 10\\",
        \\"IdNaselje\\": 110000,
        \\"Telefon\\": \\"0607654321\\"
      },
      \\"TezinaPosiljke\\": 1000,
      \\"VrednostPosiljke\\": 500000,
      \\"BrojOtkupnice\\": \\"123456\\",
      \\"Sadrzaj\\": \\"Test package\\"
    }],
    \\"DatumPrijema\\": \\"2025-01-08\\"
  }"
}'''
    add_code_block(doc, json2, "json")
    
    doc.add_paragraph('Odgovor je identican - "Korisnicko ime TEST nije registrovano!"')
    doc.add_paragraph()
    
    # Section 4: STA SMO PROVERILI
    doc.add_heading('STA SMO PROVERILI', level=2)
    
    doc.add_heading('Ispravnost polja', level=3)
    doc.add_paragraph('- Koristimo IdVrstaTranskacije sa slovom "k" (ne "c")')
    doc.add_paragraph('- Sva polja sa velikim slovom kao u dokumentaciji')
    doc.add_paragraph('- Servis = 3 za B2B')
    doc.add_paragraph('- TipSerijalizacije = 2 za JSON')
    doc.add_paragraph()
    
    doc.add_heading('Razlicite varijante autorizacije', level=3)
    doc.add_paragraph('Testirali smo nekoliko varijanti formiranja Klijent objekta:')
    doc.add_paragraph()
    
    doc.add_paragraph('1. Varijanta iz dokumentacije:')
    add_code_block(doc, '{"Username":"TEST","Password":"t3st","Jezik":"LAT","IdTipUredjaja":2}', "json")
    
    doc.add_paragraph('2. Sa malim slovima:')
    add_code_block(doc, '{"username":"TEST","password":"t3st","jezik":"LAT","idTipUredjaja":2}', "json")
    
    doc.add_paragraph('3. Razlicite kombinacije velikih/malih slova za Username i Password')
    doc.add_paragraph()
    doc.add_paragraph('Rezultat: Sve varijante vracaju "Korisnicko ime TEST nije registrovano!"')
    doc.add_paragraph()
    
    # Section 5: LOGOVI KOMPLETNOG TESTA
    doc.add_heading('LOGOVI KOMPLETNOG TESTA', level=2)
    logs = '''$ go run backend/scripts/test_wsp_minimal.go

========================================
  MINIMAL WSP API TEST
========================================

1. Testing IdVrstaTransakcije (with 'c')
Request: {"IdTransakcija":"test-1757341551","IdVrstaTransakcije":3,...}
Status: 200
Error message: Nepoznata vrsta transakcije (NapraviObjIn)! IdVrstaTransakcije = 0
Failed with Rezultat=3

2. Testing IdVrstaTranskacije (with 'k')
Request: {"IdTransakcija":"test-1757341552","IdVrstaTranskacije":3,...}
Status: 200
Error message: Korisnicko ime TEST nije registrovano!
Failed with Rezultat=3'''
    add_code_block(doc, logs, "bash")
    doc.add_paragraph()
    
    # Section 6: KLJUCNO ZAPAZANJE
    doc.add_heading('KLJUCNO ZAPAZANJE', level=2)
    doc.add_paragraph('Kada koristimo pogresno polje IdVrstaTransakcije (sa "c"), dobijamo:')
    add_code_block(doc, '"Nepoznata vrsta transakcije (NapraviObjIn)! IdVrstaTransakcije = 0"', "")
    
    doc.add_paragraph('Kada koristimo ispravno polje IdVrstaTranskacije (sa "k"), dobijamo:')
    add_code_block(doc, '"Korisnicko ime TEST nije registrovano!"', "")
    
    doc.add_paragraph('Ovo dokazuje da:')
    doc.add_paragraph('1. Nas zahtev je pravilno strukturiran')
    doc.add_paragraph('2. API ispravno parsira nas zahtev')
    doc.add_paragraph('3. Transakcija se prepoznaje korektno')
    doc.add_paragraph('4. Pristupni podaci TEST/t3st nisu aktivni ili ne postoje')
    doc.add_paragraph()
    
    # Section 7: POREDJENJE SA DOKUMENTACIJOM
    doc.add_heading('POREDJENJE SA DOKUMENTACIJOM', level=2)
    doc.add_paragraph('Iz vase dokumentacije (https://www.posta.rs/wsp-help/uvod/uvod.aspx):')
    doc.add_paragraph('- Username: TEST')
    doc.add_paragraph('- Password: t3st')
    doc.add_paragraph('- Test URL: http://212.62.32.201/WspWebApi/transakcija')
    doc.add_paragraph()
    doc.add_paragraph('Sta mi koristimo:')
    doc.add_paragraph('- Username: TEST (tacno poklapanje)')
    doc.add_paragraph('- Password: t3st (tacno poklapanje)')
    doc.add_paragraph('- URL: http://212.62.32.201/WspWebApi/transakcija (tacno poklapanje)')
    doc.add_paragraph()
    
    # Section 8: PITANJA
    doc.add_heading('PITANJA', level=2)
    doc.add_paragraph('1. Da li su pristupni podaci TEST/t3st aktivni u test okruzenju?')
    doc.add_paragraph('   - Mozda su deaktivirani?')
    doc.add_paragraph('   - Da li je potrebna prethodna registracija?')
    doc.add_paragraph()
    doc.add_paragraph('2. Postoje li dodatni zahtevi za aktivaciju?')
    doc.add_paragraph('   - IP whitelist?')
    doc.add_paragraph('   - Prethodna registracija preko emaila?')
    doc.add_paragraph('   - Specijalni header-i u zahtevu?')
    doc.add_paragraph()
    doc.add_paragraph('3. Mozete li obezbediti radni primer zahteva?')
    doc.add_paragraph('   - Sa aktivnim test pristupnim podacima')
    doc.add_paragraph('   - Koji uspesno kreira posiljku')
    doc.add_paragraph()
    
    # Add page break before appendix
    doc.add_page_break()
    
    # PRILOZI
    doc.add_heading('PRILOZI', level=2)
    
    doc.add_heading('Kompletan test skript (Go)', level=3)
    full_script = '''package main

import (
    "bytes"
    "encoding/json"
    "fmt"
    "io"
    "net/http"
    "time"
)

func main() {
    endpoint := "http://212.62.32.201/WspWebApi/transakcija"
    
    request := map[string]interface{}{
        "StrKlijent":         `{"Username":"TEST","Password":"t3st","Jezik":"LAT","IdTipUredjaja":2}`,
        "Servis":             3,
        "IdVrstaTranskacije": 3,
        "TipSerijalizacije":  2,
        "IdTransakcija":      fmt.Sprintf("test-%d", time.Now().Unix()),
        "StrIn":              `{"Naziv":"Novi"}`,
    }
    
    jsonData, _ := json.Marshal(request)
    fmt.Printf("Request: %s\\n", string(jsonData))
    
    req, _ := http.NewRequest("POST", endpoint, bytes.NewBuffer(jsonData))
    req.Header.Set("Content-Type", "application/json")
    
    client := &http.Client{Timeout: 10 * time.Second}
    resp, _ := client.Do(req)
    defer resp.Body.Close()
    
    body, _ := io.ReadAll(resp.Body)
    fmt.Printf("Response: %s\\n", string(body))
}'''
    add_code_block(doc, full_script, "go")
    
    doc.add_heading('Komanda za brzi test', level=3)
    curl_test = '''curl -X POST http://212.62.32.201/WspWebApi/transakcija \\
  -H "Content-Type: application/json" \\
  -d '{"StrKlijent":"{\\"Username\\":\\"TEST\\",\\"Password\\":\\"t3st\\",\\"Jezik\\":\\"LAT\\",\\"IdTipUredjaja\\":2}","Servis":3,"IdVrstaTranskacije":3,"TipSerijalizacije":2,"IdTransakcija":"test-123","StrIn":"{\\"Naziv\\":\\"Novi\\"}"}'
'''
    add_code_block(doc, curl_test, "bash")
    
    # Add closing note
    doc.add_paragraph()
    doc.add_paragraph('---')
    doc.add_paragraph()
    doc.add_paragraph(
        'P.S. Sigurni smo da drugim klijentima integracija radi uspesno. '
        'Molimo vas da nam pomognete da pronadjemo sta radimo pogresno. '
        'Mozda postoji neki neocigledni korak koji propustamo?'
    )
    
    # Save document
    output_path = '/data/hostel-booking-system/docs/post-express/POST_EXPRESS_TECHNICAL_REPORT.docx'
    doc.save(output_path)
    print(f"Document saved to: {output_path}")
    return output_path

if __name__ == "__main__":
    create_technical_report()