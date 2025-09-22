#!/usr/bin/env python3
"""
Create DOCX version of partner meeting one-pager
"""

import subprocess
import sys
import os

# Try to import python-docx, install if not available
try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
    from docx.enum.style import WD_STYLE_TYPE
except ImportError:
    print("Installing python-docx...")
    subprocess.check_call([sys.executable, "-m", "pip", "install", "--user", "python-docx"])
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
    from docx.enum.style import WD_STYLE_TYPE

# Create document
doc = Document()

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(11)

# Title
title = doc.add_heading('SVE TU PLATFORMA - ПАРТНЕРСКОЕ ПРЕДЛОЖЕНИЕ', 0)
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Section: СУТЬ ПРОЕКТА
doc.add_heading('🎯 СУТЬ ПРОЕКТА', level=1)
p = doc.add_paragraph()
p.add_run('Маркетплейс нового поколения для Сербии').bold = True
p.add_run(' с AI-технологиями. MVP готов на 95%, запуск октябрь 2025.')
doc.add_paragraph('• Уникальность: Загрузил фото → AI создал объявление → продал товар')
doc.add_paragraph('• Рынок: 2M+ потенциальных пользователей, €2B+ оборот C2C/B2C в Сербии')
doc.add_paragraph('• Статус: 700K строк кода, 10 месяцев разработки, €212K уже инвестировано')

# Section: ГЛАВНЫЕ ВЫЗОВЫ
doc.add_heading('🔴 ГЛАВНЫЕ ВЫЗОВЫ (где нужна помощь)', level=1)

doc.add_heading('1. ЭКВАЙРИНГ (критический блокер)', level=2)
doc.add_paragraph('• Проблема: Банки требуют историю, PaySpot тянет с подключением')
doc.add_paragraph('• Варианты: Stripe через EU компанию / Местные банки / Крипто-шлюз')
doc.add_paragraph('• Нужно: Контакты в банках или опыт получения эквайринга в Сербии')

doc.add_heading('2. ПЕРВЫЕ 1000 ПРОДАВЦОВ', level=2)
doc.add_paragraph('• Проблема: KupujemProdajem доминирует 20+ лет, люди привыкли')
doc.add_paragraph('• Решение: Оффлайн присутствие на рынках, door-to-door, личные встречи')
doc.add_paragraph('• Нужно: Опыт guerrilla маркетинга и построения комьюнити с нуля')

doc.add_heading('3. B2B ПРОДАЖИ', level=2)
doc.add_paragraph('• Цель: 200 магазинов за 3 месяца')
doc.add_paragraph('• Проблема: Консервативность локального бизнеса')
doc.add_paragraph('• Нужно: Сеть контактов в ритейле или опыт B2B продаж')

# Section: ЭКОНОМИКА И ПОТЕНЦИАЛ
doc.add_heading('💰 ЭКОНОМИКА И ПОТЕНЦИАЛ', level=1)

doc.add_heading('Unit-экономика работает:', level=2)
doc.add_paragraph('• CAC: €5-10 (оффлайн), €15-20 (онлайн)')
doc.add_paragraph('• LTV: €50-150 на пользователя')
doc.add_paragraph('• Маржа на AI: 97% (себестоимость €0.025, продажа €0.85)')
doc.add_paragraph('• Комиссии: 3-5% C2C, 8-12% B2C')

doc.add_heading('Реалистичный прогноз:', level=2)
doc.add_paragraph('• Год 1: 10K пользователей, €100K выручка (органический рост)')
doc.add_paragraph('• Год 2: 50K пользователей, €500K выручка')
doc.add_paragraph('• Год 3: 200K пользователей, €2M выручка')
doc.add_paragraph('• Exit: €10-20M через 3-5 лет (продажа региональному игроку)')

# Section: ПЛАН ДЕЙСТВИЙ
doc.add_heading('🎬 ПЛАН ДЕЙСТВИЙ НА 3 МЕСЯЦА', level=1)

doc.add_heading('Октябрь 2025 - ЗАПУСК', level=2)
doc.add_paragraph('• Soft launch с 100 друзьями')
doc.add_paragraph('• Launch party (200 гостей)')
doc.add_paragraph('• PR в локальных медиа')

doc.add_heading('Ноябрь 2025 - НОВИ САД', level=2)
doc.add_paragraph('• Захват всех рынков города')
doc.add_paragraph('• 100 B2B договоров')
doc.add_paragraph('• 1000 пользователей')

doc.add_heading('Декабрь 2025 - НОВОГОДНИЙ БУМ', level=2)
doc.add_paragraph('• Ярмарки и подарки')
doc.add_paragraph('• 5000 пользователей')
doc.add_paragraph('• €50K GMV')

# Section: ЧТО ПРЕДЛАГАЕМ ПАРТНЕРУ
doc.add_heading('🤝 ЧТО ПРЕДЛАГАЕМ ПАРТНЕРУ', level=1)

doc.add_heading('Варианты участия:', level=2)

doc.add_heading('A. Операционный партнер (CEO/COO)', level=3)
doc.add_paragraph('• Доля: 15-25% с vesting 4 года')
doc.add_paragraph('• Фокус: Продажи, партнерства, операции')
doc.add_paragraph('• Требования: Full-time с января 2026')

doc.add_heading('B. Стратегический советник', level=3)
doc.add_paragraph('• Доля: 3-5% или fee €2-5K/месяц')
doc.add_paragraph('• Время: 20-30 часов/месяц')
doc.add_paragraph('• Фокус: Стратегия, контакты, fundraising')

doc.add_heading('C. Инвестор + Партнер', level=3)
doc.add_paragraph('• Инвестиция: €20-50K')
doc.add_paragraph('• Доля: 10-20%')
doc.add_paragraph('• Роль: Board member + активное участие')

# Section: НАШИ СИЛЬНЫЕ СТОРОНЫ
doc.add_heading('💪 НАШИ СИЛЬНЫЕ СТОРОНЫ', level=1)
doc.add_paragraph('✅ Продукт готов - не идея, а работающий MVP')
doc.add_paragraph('✅ AI-преимущество - уникально для региона')
doc.add_paragraph('✅ Техническая команда - senior разработчики')
doc.add_paragraph('✅ Знание рынка - живем в Сербии 4+ года')
doc.add_paragraph('✅ Низкий burn rate - €5K/месяц')

# Section: ЧЕСТНО О РИСКАХ
doc.add_heading('⚠️ ЧЕСТНО О РИСКАХ', level=1)
doc.add_paragraph('❌ Сильная конкуренция (KupujemProdajem)')
doc.add_paragraph('❌ Нет эквайринга (пока)')
doc.add_paragraph('❌ Ограниченный бюджет на маркетинг')
doc.add_paragraph('❌ Команда без опыта в e-commerce')
doc.add_paragraph('❌ Консервативный рынок')

# Section: ТРИ СЦЕНАРИЯ РАЗВИТИЯ
doc.add_heading('🎯 ТРИ СЦЕНАРИЯ РАЗВИТИЯ', level=1)

doc.add_heading('1. ОРГАНИЧЕСКИЙ (без инвестиций)', level=2)
doc.add_paragraph('• Рост 20-30% в месяц')
doc.add_paragraph('• Самоокупаемость через 12 месяцев')
doc.add_paragraph('• €1M выручка к году 3')

doc.add_heading('2. С ИНВЕСТИЦИЯМИ (€100-250K)', level=2)
doc.add_paragraph('• Агрессивный маркетинг')
doc.add_paragraph('• Быстрый захват рынка')
doc.add_paragraph('• €5M выручка к году 3')

doc.add_heading('3. ПРОДАЖА ТЕХНОЛОГИИ', level=2)
doc.add_paragraph('• White label для других стран')
doc.add_paragraph('• Лицензирование AI-модуля')
doc.add_paragraph('• €50-100K/год пассивного дохода')

# Section: КОНКРЕТНЫЕ ВОПРОСЫ
doc.add_heading('📞 КОНКРЕТНЫЕ ВОПРОСЫ ДЛЯ ОБСУЖДЕНИЯ', level=1)
doc.add_paragraph('1. Как бы ты решил проблему эквайринга?')
doc.add_paragraph('2. Какой канал привлечения первых пользователей сработает?')
doc.add_paragraph('3. Готов ли ты к операционной роли или больше advisory?')
doc.add_paragraph('4. Есть ли у тебя контакты в:')
doc.add_paragraph('   • Сербских банках?')
doc.add_paragraph('   • Ритейл сетях (IDEA, Maxi, Roda)?')
doc.add_paragraph('   • Медиа (RTS, Blic, Kurir)?')
doc.add_paragraph('   • IT сообществе?')
doc.add_paragraph('5. Какую долю/компенсацию считаешь справедливой?')

# Section: ПОЧЕМУ СЕЙЧАС
doc.add_heading('⏰ ПОЧЕМУ СЕЙЧАС?', level=1)
doc.add_paragraph('• Окно возможностей - KupujemProdajem устарел, но еще доминирует')
doc.add_paragraph('• AI-революция - первые, кто внедряет в регионе')
doc.add_paragraph('• Команда готова - продукт есть, нужен бизнес-лидер')
doc.add_paragraph('• Низкая конкуренция - новые игроки еще не пришли')

# Section: NEXT STEPS
doc.add_heading('NEXT STEPS:', level=1)
doc.add_paragraph('1. Неформальное обсуждение за кофе')
doc.add_paragraph('2. Демо платформы (15 минут)')
doc.add_paragraph('3. Встреча с технической командой')
doc.add_paragraph('4. Решение о формате сотрудничества')
doc.add_paragraph('5. Начало работы - октябрь/ноябрь 2025')

doc.add_paragraph()
doc.add_paragraph('Demo: dev.svetu.rs')
doc.add_paragraph('Pitch deck: можем выслать после встречи')

# Save document
output_path = '/data/hostel-booking-system/docs/SVE_TU_PARTNER_PROPOSAL.docx'
doc.save(output_path)
print(f"Document saved to: {output_path}")